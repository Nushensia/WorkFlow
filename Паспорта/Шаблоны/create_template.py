#!/usr/bin/env python3
"""
Создание DOCX-шаблона протокола физико-химического фазового анализа
с рамкой (таблица-обёртка) и нижним колонтитулом (footer с PAGE/NUMPAGES).
"""

import os, zipfile, copy, shutil
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUTPUT = os.path.join(os.path.dirname(__file__), "Шаблон_протокола_фазового_анализа.docx")

# ═══════════════════════════════════════════════
# 1. Создаём документ с контентом (без рамки и footer)
# ═══════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)

def add_centered(text, bold=False, size=12, space_after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_justified(text, bold=False, space_after=0, size=12, italic=False, left_indent=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_left(text, bold=False, size=12, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_mixed_field(label, placeholder, bold_label=True, bold_ph=False, italic_ph=False):
    """Поле с жирной меткой и обычным плейсхолдером."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    r1.bold = bold_label
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(placeholder)
    r2.bold = bold_ph
    r2.italic = italic_ph
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p

def set_cell_text(cell, text, bold=False, size=9, align='center'):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def shade_cell(cell, color='D9E2F3'):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

# ── Собираем контент в список для вставки в строку 1 рамки ──
# Будем добавлять прямо в документ, а потом через lxml переместим

# Заголовок (будет внутри строки 1 рамки)
add_centered('ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА', bold=True, size=14, space_after=6)

# Номер и дата
p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_num.paragraph_format.line_spacing = 1.5
p_num.paragraph_format.space_before = Pt(0)
p_num.paragraph_format.space_after = Pt(12)
r = p_num.add_run('№ [номер]/[индекс]-[год] от «[дд]» [месяц] [гггг] г.')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

# Поля 1-8
add_mixed_field('1.\tНаименование Заказчика испытаний: ', '[заказчик]')
add_mixed_field('2.\tНаименование конкретной испытуемой продукции: ', '[наименование_продукции]')
add_mixed_field('3.\tКраткая характеристика испытуемого образца: ', '[марка_стали]')
add_mixed_field('4.\tВид испытаний: ', 'Физико-химический фазовый анализ.')
add_mixed_field('5.\tНормативные документы, использованные при испытаниях, в т. ч. методики: ', '[нормативные_документы]')
add_mixed_field('6.\tИспытательное оборудование: ', '[оборудование]')
add_mixed_field('7.\tКоличество испытанных образцов и даты начала и окончания проведения испытаний: ', '[образцы_даты]')

# п.8
p8 = doc.add_paragraph()
p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p8.paragraph_format.line_spacing = 1.5
p8.paragraph_format.space_before = Pt(0)
p8.paragraph_format.space_after = Pt(0)
r = p8.add_run('8.\tРезультаты испытаний:')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_justified('[условия_электрохимического_травления]', bold=False, space_after=6, italic=True)
add_justified('[режимы_рентгенографирования]', bold=False, space_after=12, italic=True)
add_justified('[описание_фазового_состава]', bold=False, space_after=6, italic=True)

# Подпись таблицы
add_centered('Таблица 1 — Фазовый состав образца [обозначение]', bold=False, size=11, space_after=6)

# Таблица фаз
headers = ['№', 'Phase name', 'Chemical formula', 'Crystal system', 'Space group', 'Lattice parameters, Å', 'DB card number']
table = doc.add_table(rows=2, cols=7)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers):
    align = 'center' if i != 1 else 'left'  # Phase name — left
    set_cell_text(table.cell(0, i), h, bold=True, size=9, align=align)
    shade_cell(table.cell(0, i), 'D9E2F3')

example = ['1', '[Phase name]', '[Formula]', '[Sys]', '[SG]', '[a,b,c,α,β,γ]', '[DB#]']
for i, val in enumerate(example):
    align = 'center' if i != 1 else 'left'
    set_cell_text(table.cell(1, i), val, bold=False, size=9, align=align)

widths = [Cm(0.8), Cm(3.0), Cm(2.5), Cm(2.0), Cm(2.5), Cm(3.5), Cm(2.5)]
for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

doc.add_paragraph()  # отступ

# Рисунок
add_centered('[Место для вставки рисунка — Дифрактограмма]', bold=False, size=11, space_after=6, italic=True)
add_centered('Рисунок 1 — Дифрактограмма образца [обозначение]', bold=False, size=11, space_after=12)

# Заключение
add_justified('Заключение:', bold=True, space_after=6)
for item in ['• [карбиды]', '• [оксиды]', '• [прочие фазы]']:
    add_justified(item, bold=False, space_after=0, left_indent=1.0)

doc.add_paragraph()

# Примечание
p_note_label = doc.add_paragraph()
p_note_label.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_note_label.paragraph_format.line_spacing = 1.5
p_note_label.paragraph_format.space_before = Pt(0)
p_note_label.paragraph_format.space_after = Pt(0)
r = p_note_label.add_run('Примечание:')
r.bold = True
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

add_justified('[текст_примечания]', bold=False, space_after=12, italic=True)

# Подпись
doc.add_paragraph()
add_justified('Заместитель технического директора по развитию и науке', bold=False, space_after=0)
doc.add_paragraph()
add_justified('___________________________  [подпись]', bold=False, space_after=0)
add_justified('[ФИО]', bold=False, space_after=0)

# ── Сохраняем через python-docx ──
TMP = OUTPUT.replace('.docx', '_tmp.docx')
doc.save(TMP)

# ═══════════════════════════════════════════════
# 2. Через lxml + zipfile: добавляем рамку, шапку, footer
# ═══════════════════════════════════════════════
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
ns_rel = 'http://schemas.openxmlformats.org/package/2006/relationships'
ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'

W = f'{{{ns_w}}}'
R = f'{{{ns_r}}}'

def mk(tag):
    """Create an OxmlElement with namespace."""
    el = OxmlElement(tag)
    return el

def mk_attr(name, value):
    """Helper to assign w: prefixed attributes."""
    return {f'{W}{name}': value}

# Открываем сохранённый временный файл
with zipfile.ZipFile(TMP, 'r') as z:
    data = {n: z.read(n) for n in z.namelist()}

# Парсим document.xml
doc_xml = etree.fromstring(data['word/document.xml'])
body = doc_xml.find(f'{W}body')

# Находим sectPr (не удаляем из дерева!)
sectPr = body.find(f'{W}sectPr')

# Сохраняем все дочерние элементы body, которые идут ДО sectPr — это контент для строки 1
content_elements = []
found_sectpr = False
for child in list(body):
    if child is sectPr:
        found_sectpr = True
        break
    content_elements.append(child)

# Удаляем контент из body (sectPr остаётся!)
for el in content_elements:
    body.remove(el)

# ═══════════════════════════════════════════════
# Создаём таблицу-рамку (1 колонка, 2 строки)
# ═══════════════════════════════════════════════
frame_tbl = OxmlElement('w:tbl')

# tblPr
tblPr = OxmlElement('w:tblPr')
tblW = OxmlElement('w:tblW')
tblW.set(f'{W}w', '0')
tblW.set(f'{W}type', 'auto')
tblPr.append(tblW)

# Borders
tblBorders = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(f'{W}val', 'single')
    b.set(f'{W}sz', '18')
    b.set(f'{W}space', '0')
    b.set(f'{W}color', 'auto')
    tblBorders.append(b)
tblPr.append(tblBorders)
tblLook = OxmlElement('w:tblLook')
tblLook.set(f'{W}val', '04A0')
tblPr.append(tblLook)
frame_tbl.append(tblPr)

# tblGrid — 1 колонка
tblGrid = OxmlElement('w:tblGrid')
gridCol = OxmlElement('w:gridCol')
gridCol.set(f'{W}w', '9309')
tblGrid.append(gridCol)
frame_tbl.append(tblGrid)

# ── Строка 0: Шапка организации ──
tr0 = OxmlElement('w:tr')
tc0 = OxmlElement('w:tc')
tcPr0 = OxmlElement('w:tcPr')
tcW0 = OxmlElement('w:tcW')
tcW0.set(f'{W}w', '10139')
tcW0.set(f'{W}type', 'dxa')
tcPr0.append(tcW0)
tcBorders0 = OxmlElement('w:tcBorders')
for side in ['top', 'left', 'bottom', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(f'{W}val', 'single')
    b.set(f'{W}sz', '18')
    b.set(f'{W}space', '0')
    b.set(f'{W}color', 'auto')
    tcBorders0.append(b)
tcPr0.append(tcBorders0)
hideMark = OxmlElement('w:hideMark')
tcPr0.append(hideMark)
tc0.append(tcPr0)

# Вложенная таблица шапки (2 колонки, без границ)
nested_tbl = OxmlElement('w:tbl')

# tblPr nested
ntblPr = OxmlElement('w:tblPr')
ntblW = OxmlElement('w:tblW')
ntblW.set(f'{W}w', '0')
ntblW.set(f'{W}type', 'auto')
ntblPr.append(ntblW)
tblInd = OxmlElement('w:tblInd')
tblInd.set(f'{W}w', '108')
tblInd.set(f'{W}type', 'dxa')
ntblPr.append(tblInd)
ntblBorders = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement(f'w:{side}')
    b.set(f'{W}val', 'none')
    b.set(f'{W}sz', '0')
    b.set(f'{W}space', '0')
    b.set(f'{W}color', 'auto')
    ntblBorders.append(b)
ntblPr.append(ntblBorders)
ntblLook = OxmlElement('w:tblLook')
ntblLook.set(f'{W}val', '04A0')
ntblPr.append(ntblLook)

# tblGrid nested — 1 колонка на всю ширину
ntblGrid = OxmlElement('w:tblGrid')
ngc1 = OxmlElement('w:gridCol')
ngc1.set(f'{W}w', '9309')
ntblGrid.append(ngc1)

# Строка 0 вложенной таблицы — логотип
ntr0 = OxmlElement('w:tr')
ntc0 = OxmlElement('w:tc')
ntcPr0 = OxmlElement('w:tcPr')
ntcW0 = OxmlElement('w:tcW')
ntcW0.set(f'{W}w', '9309')
ntcW0.set(f'{W}type', 'dxa')
ntcPr0.append(ntcW0)
ntc0.append(ntcPr0)
np0 = OxmlElement('w:p')
ntc0.append(np0)
ntr0.append(ntc0)

# Строка 1 вложенной таблицы — информация о компании
ntr1 = OxmlElement('w:tr')
header_texts = [
    ('Акционерное общество «НПО «Ленкор»', True),
    ('Лаборатория разрушающих методов контроля', False),
    ('Свидетельство об аккредитации', False),
    ('ОАО «НТЦ «Промышленная безопасность» № ИЛ/ЛРИ-02436', False),
    ('Адрес: Россия, 192236, г. Санкт-Петербург, ул. Белы Куна, д. 30, литера А, пом. 25-Н, офис 1408 тел. 8(812) 335-13-27; E-mail: office@npo-lencor.ru', False),
]

ntc1 = OxmlElement('w:tc')
ntcPr1 = OxmlElement('w:tcPr')
ntcW1 = OxmlElement('w:tcW')
ntcW1.set(f'{W}w', '9309')
ntcW1.set(f'{W}type', 'dxa')
ntcPr1.append(ntcW1)
vAlign1 = OxmlElement('w:vAlign')
vAlign1.set(f'{W}val', 'center')
ntcPr1.append(vAlign1)
ntc1.append(ntcPr1)

for txt, bold in header_texts:
    np1 = OxmlElement('w:p')
    npPr1 = OxmlElement('w:pPr')
    npJc1 = OxmlElement('w:jc')
    npJc1.set(f'{W}val', 'center')
    npPr1.append(npJc1)
    npSpacing1 = OxmlElement('w:spacing')
    npSpacing1.set(f'{W}line', '240')
    npSpacing1.set(f'{W}lineRule', 'auto')
    npSpacing1.set(f'{W}before', '0')
    npSpacing1.set(f'{W}after', '0')
    npPr1.append(npSpacing1)
    np1.append(npPr1)
    
    nr1 = OxmlElement('w:r')
    nrPr1 = OxmlElement('w:rPr')
    nrf1 = OxmlElement('w:rFonts')
    nrf1.set(f'{W}ascii', 'Times New Roman')
    nrf1.set(f'{W}hAnsi', 'Times New Roman')
    nrPr1.append(nrf1)
    nrsz1 = OxmlElement('w:sz')
    nrsz1.set(f'{W}val', '24')
    nrPr1.append(nrsz1)
    if bold:
        nb = OxmlElement('w:b')
        nrPr1.append(nb)
    if txt.startswith('Адрес:'):
        nrsz1.set(f'{W}val', '22')
    nr1.append(nrPr1)
    nt1 = OxmlElement('w:t')
    nt1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    nt1.text = txt
    nr1.append(nt1)
    np1.append(nr1)
    ntc1.append(np1)

ntr1.append(ntc1)

ntblPr.append(ntblGrid)
nested_tbl.append(ntblPr)
nested_tbl.append(ntr0)
nested_tbl.append(ntr1)

tc0.append(nested_tbl)
tr0.append(tc0)
frame_tbl.append(tr0)

# ── Строка 1: Контент ──
tr1 = OxmlElement('w:tr')
tc1 = OxmlElement('w:tc')
tcPr1 = OxmlElement('w:tcPr')
tcW1 = OxmlElement('w:tcW')
tcW1.set(f'{W}w', '10139')
tcW1.set(f'{W}type', 'dxa')
tcPr1.append(tcW1)
tcBorders1 = OxmlElement('w:tcBorders')
for side in ['top', 'left', 'bottom', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(f'{W}val', 'single')
    b.set(f'{W}sz', '18')
    b.set(f'{W}space', '0')
    b.set(f'{W}color', 'auto')
    tcBorders1.append(b)
tcPr1.append(tcBorders1)
tc1.append(tcPr1)

# Переносим весь контент
for el in content_elements:
    tc1.append(el)

tr1.append(tc1)
frame_tbl.append(tr1)

# Вставляем таблицу-рамку в body
body.insert(0, frame_tbl)

# sectPr уже в дереве — ничего не делаем





# ═══════════════════════════════════════════════
# 3. Создаём footer с PAGE/NUMPAGES
# ═══════════════════════════════════════════════
footer_xml = etree.Element(f'{W}ftr')

footer_para = etree.SubElement(footer_xml, f'{W}p')
fpPr = etree.SubElement(footer_para, f'{W}pPr')
fjc = etree.SubElement(fpPr, f'{W}jc')
fjc.set(f'{W}val', 'center')
fspacing = etree.SubElement(fpPr, f'{W}spacing')
fspacing.set(f'{W}line', '240')
fspacing.set(f'{W}lineRule', 'auto')
fspacing.set(f'{W}before', '0')
fspacing.set(f'{W}after', '0')

# Часть 1: "ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА № [номер]/[индекс]-[год]   Лист "
fr1 = etree.SubElement(footer_para, f'{W}r')
frPr1 = etree.SubElement(fr1, f'{W}rPr')
ff1 = etree.SubElement(frPr1, f'{W}rFonts')
ff1.set(f'{W}ascii', 'Times New Roman')
ff1.set(f'{W}hAnsi', 'Times New Roman')
fsz1 = etree.SubElement(frPr1, f'{W}sz')
fsz1.set(f'{W}val', '20')  # 10pt
ft1 = etree.SubElement(fr1, f'{W}t')
ft1.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
ft1.text = 'ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА № [номер]/[индекс]-[год]   Лист '

# Часть 2: PAGE (поле, полужирное)
fr2_begin = etree.SubElement(footer_para, f'{W}r')
frPr2 = etree.SubElement(fr2_begin, f'{W}rPr')
fb2 = etree.SubElement(frPr2, f'{W}b')
fbCs2 = etree.SubElement(frPr2, f'{W}bCs')
ff2 = etree.SubElement(frPr2, f'{W}rFonts')
ff2.set(f'{W}ascii', 'Times New Roman')
ff2.set(f'{W}hAnsi', 'Times New Roman')
fsz2 = etree.SubElement(frPr2, f'{W}sz')
fsz2.set(f'{W}val', '20')
ffld1 = etree.SubElement(fr2_begin, f'{W}fldChar')
ffld1.set(f'{W}fldCharType', 'begin')

fr2_instr = etree.SubElement(footer_para, f'{W}r')
frPr2i = etree.SubElement(fr2_instr, f'{W}rPr')
fb2i = etree.SubElement(frPr2i, f'{W}b')
fbCs2i = etree.SubElement(frPr2i, f'{W}bCs')
ff2i = etree.SubElement(frPr2i, f'{W}rFonts')
ff2i.set(f'{W}ascii', 'Times New Roman')
ff2i.set(f'{W}hAnsi', 'Times New Roman')
fsz2i = etree.SubElement(frPr2i, f'{W}sz')
fsz2i.set(f'{W}val', '20')
finstr = etree.SubElement(fr2_instr, f'{W}instrText')
finstr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
finstr.text = 'PAGE'

fr2_sep = etree.SubElement(footer_para, f'{W}r')
frPr2s = etree.SubElement(fr2_sep, f'{W}rPr')
fb2s = etree.SubElement(frPr2s, f'{W}b')
fbCs2s = etree.SubElement(frPr2s, f'{W}bCs')
ff2s = etree.SubElement(frPr2s, f'{W}rFonts')
ff2s.set(f'{W}ascii', 'Times New Roman')
ff2s.set(f'{W}hAnsi', 'Times New Roman')
fsz2s = etree.SubElement(frPr2s, f'{W}sz')
fsz2s.set(f'{W}val', '20')
ffld2 = etree.SubElement(fr2_sep, f'{W}fldChar')
ffld2.set(f'{W}fldCharType', 'separate')

fr2_text = etree.SubElement(footer_para, f'{W}r')
frPr2t = etree.SubElement(fr2_text, f'{W}rPr')
fb2t = etree.SubElement(frPr2t, f'{W}b')
fbCs2t = etree.SubElement(frPr2t, f'{W}bCs')
ff2t = etree.SubElement(frPr2t, f'{W}rFonts')
ff2t.set(f'{W}ascii', 'Times New Roman')
ff2t.set(f'{W}hAnsi', 'Times New Roman')
fsz2t = etree.SubElement(frPr2t, f'{W}sz')
fsz2t.set(f'{W}val', '20')
ft2 = etree.SubElement(fr2_text, f'{W}t')
ft2.text = '1'

fr2_end = etree.SubElement(footer_para, f'{W}r')
frPr2e = etree.SubElement(fr2_end, f'{W}rPr')
fb2e = etree.SubElement(frPr2e, f'{W}b')
fbCs2e = etree.SubElement(frPr2e, f'{W}bCs')
ff2e = etree.SubElement(frPr2e, f'{W}rFonts')
ff2e.set(f'{W}ascii', 'Times New Roman')
ff2e.set(f'{W}hAnsi', 'Times New Roman')
fsz2e = etree.SubElement(frPr2e, f'{W}sz')
fsz2e.set(f'{W}val', '20')
ffld3 = etree.SubElement(fr2_end, f'{W}fldChar')
ffld3.set(f'{W}fldCharType', 'end')

# Часть 3: " из "
fr3 = etree.SubElement(footer_para, f'{W}r')
frPr3 = etree.SubElement(fr3, f'{W}rPr')
ff3 = etree.SubElement(frPr3, f'{W}rFonts')
ff3.set(f'{W}ascii', 'Times New Roman')
ff3.set(f'{W}hAnsi', 'Times New Roman')
fsz3 = etree.SubElement(frPr3, f'{W}sz')
fsz3.set(f'{W}val', '20')
ft3 = etree.SubElement(fr3, f'{W}t')
ft3.text = ' из '

# Часть 4: NUMPAGES (поле, полужирное)
fr4_begin = etree.SubElement(footer_para, f'{W}r')
frPr4 = etree.SubElement(fr4_begin, f'{W}rPr')
fb4 = etree.SubElement(frPr4, f'{W}b')
fbCs4 = etree.SubElement(frPr4, f'{W}bCs')
ff4 = etree.SubElement(frPr4, f'{W}rFonts')
ff4.set(f'{W}ascii', 'Times New Roman')
ff4.set(f'{W}hAnsi', 'Times New Roman')
fsz4 = etree.SubElement(frPr4, f'{W}sz')
fsz4.set(f'{W}val', '20')
ffld4 = etree.SubElement(fr4_begin, f'{W}fldChar')
ffld4.set(f'{W}fldCharType', 'begin')

fr4_instr = etree.SubElement(footer_para, f'{W}r')
frPr4i = etree.SubElement(fr4_instr, f'{W}rPr')
fb4i = etree.SubElement(frPr4i, f'{W}b')
fbCs4i = etree.SubElement(frPr4i, f'{W}bCs')
ff4i = etree.SubElement(frPr4i, f'{W}rFonts')
ff4i.set(f'{W}ascii', 'Times New Roman')
ff4i.set(f'{W}hAnsi', 'Times New Roman')
fsz4i = etree.SubElement(frPr4i, f'{W}sz')
fsz4i.set(f'{W}val', '20')
finstr2 = etree.SubElement(fr4_instr, f'{W}instrText')
finstr2.text = 'NUMPAGES'

fr4_sep = etree.SubElement(footer_para, f'{W}r')
frPr4s = etree.SubElement(fr4_sep, f'{W}rPr')
fb4s = etree.SubElement(frPr4s, f'{W}b')
fbCs4s = etree.SubElement(frPr4s, f'{W}bCs')
ff4s = etree.SubElement(frPr4s, f'{W}rFonts')
ff4s.set(f'{W}ascii', 'Times New Roman')
ff4s.set(f'{W}hAnsi', 'Times New Roman')
fsz4s = etree.SubElement(frPr4s, f'{W}sz')
fsz4s.set(f'{W}val', '20')
ffld5 = etree.SubElement(fr4_sep, f'{W}fldChar')
ffld5.set(f'{W}fldCharType', 'separate')

fr4_text = etree.SubElement(footer_para, f'{W}r')
frPr4t = etree.SubElement(fr4_text, f'{W}rPr')
fb4t = etree.SubElement(frPr4t, f'{W}b')
fbCs4t = etree.SubElement(frPr4t, f'{W}bCs')
ff4t = etree.SubElement(frPr4t, f'{W}rFonts')
ff4t.set(f'{W}ascii', 'Times New Roman')
ff4t.set(f'{W}hAnsi', 'Times New Roman')
fsz4t = etree.SubElement(frPr4t, f'{W}sz')
fsz4t.set(f'{W}val', '20')
ft4 = etree.SubElement(fr4_text, f'{W}t')
ft4.text = '1'

fr4_end = etree.SubElement(footer_para, f'{W}r')
frPr4e = etree.SubElement(fr4_end, f'{W}rPr')
fb4e = etree.SubElement(frPr4e, f'{W}b')
fbCs4e = etree.SubElement(frPr4e, f'{W}bCs')
ff4e = etree.SubElement(frPr4e, f'{W}rFonts')
ff4e.set(f'{W}ascii', 'Times New Roman')
ff4e.set(f'{W}hAnsi', 'Times New Roman')
fsz4e = etree.SubElement(frPr4e, f'{W}sz')
fsz4e.set(f'{W}val', '20')
ffld6 = etree.SubElement(fr4_end, f'{W}fldChar')
ffld6.set(f'{W}fldCharType', 'end')

data['word/footer1.xml'] = etree.tostring(footer_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════
# 4. Обновляем document.xml.rels — добавляем footer reference
# ═══════════════════════════════════════════════
rels = etree.fromstring(data['word/_rels/document.xml.rels'])
existing_ids = [int(r.get('Id', 'rId0')[3:]) for r in rels]
max_id = max(existing_ids) if existing_ids else 0

# Проверяем, есть ли уже footer rel
footer_rel_exists = any(
    'footer' in (r.get('Target', '') or '') for r in rels
)

if not footer_rel_exists:
    max_id += 1
    new_rel = etree.SubElement(rels, f'{{{ns_rel}}}Relationship')
    new_rel.set('Id', f'rId{max_id}')
    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')
    new_rel.set('Target', 'footer1.xml')
    footer_rid = f'rId{max_id}'
else:
    # Находим существующий rId для footer
    footer_rid = None
    for r in rels:
        if 'footer' in (r.get('Target', '') or ''):
            footer_rid = r.get('Id')
            break

# Обновляем sectPr в document.xml — добавляем footerReference (всегда, если отсутствует)
if sectPr is not None:
    fr = sectPr.find(f'{W}footerReference')
    if fr is None and footer_rid:
        fr_new = etree.SubElement(sectPr, f'{W}footerReference')
        fr_new.set(f'{W}type', 'default')
        fr_new.set(f'{R}id', footer_rid)
        # Перемещаем сразу после pgSz (SubElement добавляет в конец)
        pgSz = sectPr.find(f'{W}pgSz')
        if pgSz is not None:
            sectPr.remove(fr_new)
            pgSz.addnext(fr_new)
            

# Сохраняем document.xml (после добавления footerReference)
data['word/document.xml'] = etree.tostring(doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

data['word/_rels/document.xml.rels'] = etree.tostring(rels, xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════
# 5. Обновляем [Content_Types].xml
# ═══════════════════════════════════════════════
ct = etree.fromstring(data['[Content_Types].xml'])
footer_override = ct.find(f'{{{ns_ct}}}Override[@PartName="/word/footer1.xml"]')
if footer_override is None:
    ov = etree.SubElement(ct, f'{{{ns_ct}}}Override')
    ov.set('PartName', '/word/footer1.xml')
    ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml')

# Проверяем default для emf (если не было)
emf_default = ct.find(f'{{{ns_ct}}}Default[@Extension="emf"]')
if emf_default is None:
    d = etree.SubElement(ct, f'{{{ns_ct}}}Default')
    d.set('Extension', 'emf')
    d.set('ContentType', 'image/x-emf')

data['[Content_Types].xml'] = etree.tostring(ct, xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════
# 6. Сохраняем
# ═══════════════════════════════════════════════
# Удаляем старый tmp если есть
if os.path.exists(OUTPUT):
    os.remove(OUTPUT)

tmp2 = OUTPUT.replace('.docx', '_tmp2.docx')
with zipfile.ZipFile(tmp2, 'w') as z:
    for name, content in data.items():
        z.writestr(name, content)

os.replace(tmp2, OUTPUT)
os.remove(TMP)

print(f"✅ Шаблон сохранён: {OUTPUT}")
