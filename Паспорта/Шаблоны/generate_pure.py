#!/usr/bin/env python3
"""Шаблон протокола — только python-docx API, без lxml/ZIP."""

import os, shutil, tempfile, random
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from PIL import Image

DIR = r"D:\ИИ\Git Projetcs\Nushensia\WorkFlow\Паспорта\Шаблоны"
OUTPUT = os.path.join(DIR, "Шаблон_протокола_фазового_анализа.docx")
LOGO = os.path.join(DIR, "..", "Исходные протоколы", "logo_top.png")

doc = Document()

# === Формат страницы ===
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(1.5)

# === Шрифт умолчанию ===
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# === Вспомогательные функции ===
def add_para(text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             italic=False, space_before=0, space_after=0, left_indent=None,
             line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_mixed(label, placeholder, bold_label=True, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run(label)
    r1.bold = bold_label
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(size)
    if placeholder:
        r2 = p.add_run(' ' + placeholder)
        r2.bold = False
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(size)
    return p

def remove_shading(cell):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        tcPr.remove(shd)

def set_cell_width(cell, width_pct):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_pct))
    tcW.set(qn('w:type'), 'pct')

def set_cell_valign(cell, val='center'):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    va = tcPr.find(qn('w:vAlign'))
    if va is None:
        va = OxmlElement('w:vAlign')
        tcPr.append(va)
    va.set(qn('w:val'), val)

def remove_table_borders(tbl):
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.find(qn('w:' + side))
        if b is None:
            b = OxmlElement('w:' + side)
            borders.append(b)
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')

# === 1. ШАПКА (таблица-рамка 2 строки) ===
# Внешняя таблица-рамка: 1 колонка, 2 строки
frame = doc.add_table(rows=2, cols=1)
frame.alignment = WD_TABLE_ALIGNMENT.CENTER

# Настройка границ рамки (толстая черная линия 0.75pt)
tblPr = frame._tbl.find(qn('w:tblPr'))
# Удаляем стиль
tblStyle = tblPr.find(qn('w:tblStyle'))
if tblStyle is not None:
    tblPr.remove(tblStyle)

# Добавляем границы через OxmlElement
borders = OxmlElement('w:tblBorders')
for s in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    b = OxmlElement('w:' + s)
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '18')
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), 'auto')
    borders.append(b)
tblPr.append(borders)

# Ширина таблицы 100%
tw = tblPr.find(qn('w:tblW'))
if tw is None:
    tw = OxmlElement('w:tblW')
    tblPr.append(tw)
tw.set(qn('w:w'), '5000')
tw.set(qn('w:type'), 'pct')

# === Строка 0: Шапка (вложенная таблица для логотипа и инфо) ===
cell0 = frame.cell(0, 0)
remove_shading(cell0)
set_cell_valign(cell0, 'top')

# Вложенная таблица: 1 колонка × 2 строки (логотип + инфо)
header_tbl = cell0.add_table(rows=2, cols=1)
remove_table_borders(header_tbl)

# Строка 0.0: логотип
logo_cell = header_tbl.cell(0, 0)
remove_shading(logo_cell)
set_cell_valign(logo_cell, 'top')
# Добавляем пустой параграф для логотипа
if os.path.isfile(LOGO):
    img = Image.open(LOGO)
    iw, ih = img.size
    # Масштаб: ~3.5 см = 1260000 EMU
    width_emu = Emu(1260000)
    height_emu = Emu(int(1260000 * ih / iw))
    run = logo_cell.paragraphs[0].add_run()
    run.add_picture(LOGO, width=width_emu, height=height_emu)
    logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# Строка 0.1: информация о компании (по центру)
info_cell = header_tbl.cell(1, 0)
remove_shading(info_cell)
set_cell_valign(info_cell, 'center')

company_lines = [
    ('Акционерное общество «НПО «Ленкор»', True, 12),
    ('Лаборатория разрушающих методов контроля', False, 12),
    ('Свидетельство об аккредитации', False, 12),
    ('ОАО «НТЦ «Промышленная безопасность» № ИЛ/ЛРИ-02436', False, 11),
    ('Адрес: Россия, 192236, г. Санкт-Петербург, ул. Белы Куна, д. 30, литера А, пом. 25-Н, офис 1408 тел. 8(812) 335-13-27; E-mail: office@npo-lencor.ru', False, 11),
]
for text, bold, sz in company_lines:
    p = info_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(sz)
# Удаляем пустой параграф, созданный по умолчанию
for p in info_cell.paragraphs:
    if not p.text.strip():
        pp = p._element
        pp.getparent().remove(pp)

# === Строка 1: Контент ===
cell1 = frame.cell(1, 0)
remove_shading(cell1)
set_cell_valign(cell1, 'top')

# Удаляем пустой параграф по умолчанию
for p in cell1.paragraphs:
    if not p.text.strip():
        pp = p._element
        pp.getparent().remove(pp)

# Далее весь контент добавляем в cell1
# Для добавления контента в ячейку используем cell1.add_paragraph() и т.д.
content = cell1  # будем добавлять в эту ячейку

# === 2. ЗАГОЛОВОК ===
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(6)
r = p.add_run('ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

# === 3. НОМЕР И ДАТА ===
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(12)
r = p.add_run('№ [номер]/[индекс]-[год] от «[дд]» [месяц] [гггг] г.')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

# === 4. ПОЛЯ ===
fields = [
    ('1. Наименование Заказчика испытаний:', '[заказчик]'),
    ('2. Наименование конкретной испытуемой продукции:', '[наименование_продукции]'),
    ('3. Краткая характеристика испытуемого образца:', '[характеристика_образца]'),
    ('4. Вид испытаний:', '[вид_испытаний]'),
    ('5. Нормативные документы, использованные при испытаниях, в т. ч. методики:', '[нормативные_документы]'),
    ('6. Испытательное оборудование:', '[оборудование]'),
    ('7. Количество испытанных образцов и даты начала и окончания проведения испытаний:', '[количество_образцов_даты]'),
    ('8. Результаты испытаний:', None),
]
for label, placeholder in fields:
    if placeholder:
        p = content.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(label + ' ')
        r1.bold = True
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r2 = p.add_run(placeholder)
        r2.bold = False
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
    else:
        p = content.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(label)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# Условия (курсив)
for txt in ['[условия_электрохимического_травления]', '[режимы_рентгенографирования]', '[описание_фазового_состава]']:
    p = content.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(txt)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# === 5. ТАБЛИЦА ФАЗ ===
headers = ['№', 'Phase name', 'Chemical formula', 'Crystal system',
           'Space group', 'Lattice parameters, Å', 'DB card number']
col_widths_pct = [226, 765, 538, 702, 515, 1785, 470]

tbl = content.add_table(rows=2, cols=7)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Установка ширины таблицы
tbl_pr = tbl._tbl.find(qn('w:tblPr'))
tw = tbl_pr.find(qn('w:tblW'))
if tw is None:
    tw = OxmlElement('w:tblW')
    tbl_pr.append(tw)
tw.set(qn('w:w'), '5000')
tw.set(qn('w:type'), 'pct')

for i, h in enumerate(headers):
    cell = tbl.cell(0, i)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.line_spacing = 1.0
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)
    run = cp.add_run(h)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    set_cell_width(cell, col_widths_pct[i])
    remove_shading(cell)

example = ['1', '[Phase name]', '[Formula]', '[Sys]', '[SG]',
           '[a,b,c,α,β,γ]', '[DB#]']
for i, val in enumerate(example):
    cell = tbl.cell(1, i)
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.line_spacing = 1.0
    cp.paragraph_format.space_before = Pt(0)
    cp.paragraph_format.space_after = Pt(0)
    run = cp.add_run(val)
    run.bold = False
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)
    set_cell_width(cell, col_widths_pct[i])

# === 6. РИСУНОК ===
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
r = p.add_run('[Место для вставки рисунка — Дифрактограмма]')
r.italic = False
r.font.name = 'Times New Roman'
r.font.size = Pt(11)

p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Рисунок 1 — Дифрактограмма образца [обозначение]')
r.bold = False
r.font.name = 'Times New Roman'
r.font.size = Pt(11)

# === 7. ЗАКЛЮЧЕНИЕ ===
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
r = p.add_run('Заключение:')
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

for txt in ['[список_фаз]', '[количество_фаз]', '[размер_ОКР]', '[микроискажения]']:
    p = content.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(txt)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# === 8. ПРИМЕЧАНИЕ ===
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
r = p.add_run('Примечание:')
r.bold = True
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(12)
r2 = p.add_run(' [текст_примечания]')
r2.italic = True
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

# === 9. ПОДПИСЬ ===
content.add_paragraph()
p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
r = p.add_run('{{Должность}}')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = content.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = 1.5
r = p.add_run('{{ФИО}}')
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

# === FOOTER ===
footer = doc.sections[0].footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.paragraph_format.line_spacing = 1.0

# Текст
run = fp.add_run('ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА № [номер]/[индекс]-[год]   Лист ')
run.font.name = 'Times New Roman'
run.font.size = Pt(10)

# Поле PAGE
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
run1 = fp.add_run()
run1._r.append(fldChar1)

run2 = fp.add_run()
run2.bold = True
run2.font.name = 'Times New Roman'
run2.font.size = Pt(10)
instrText = OxmlElement('w:instrText')
instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
instrText.text = ' PAGE '
run2._r.append(instrText)

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
run3 = fp.add_run()
run3._r.append(fldChar2)

run4 = fp.add_run('1')
run4.bold = True
run4.font.name = 'Times New Roman'
run4.font.size = Pt(10)

fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run5 = fp.add_run()
run5._r.append(fldChar3)

# " из "
run6 = fp.add_run(' из ')
run6.font.name = 'Times New Roman'
run6.font.size = Pt(10)

# Поле NUMPAGES
fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'begin')
run7 = fp.add_run()
run7._r.append(fldChar4)

run8 = fp.add_run()
run8.bold = True
run8.font.name = 'Times New Roman'
run8.font.size = Pt(10)
instrText2 = OxmlElement('w:instrText')
instrText2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
instrText2.text = ' NUMPAGES '
run8._r.append(instrText2)

fldChar5 = OxmlElement('w:fldChar')
fldChar5.set(qn('w:fldCharType'), 'separate')
run9 = fp.add_run()
run9._r.append(fldChar5)

run10 = fp.add_run('1')
run10.bold = True
run10.font.name = 'Times New Roman'
run10.font.size = Pt(10)

fldChar6 = OxmlElement('w:fldChar')
fldChar6.set(qn('w:fldCharType'), 'end')
run11 = fp.add_run()
run11._r.append(fldChar6)

# === СОХРАНЕНИЕ ===
import zipfile as _zf
tmp = os.path.join(tempfile.gettempdir(), "hermes_clean.docx")
doc.save(tmp)

# Cleanup: remove thumbnail, fix app.xml
with _zf.ZipFile(tmp, 'r') as z:
    d = {n: z.read(n) for n in z.namelist()}

# Remove thumbnail
if 'docProps/thumbnail.jpeg' in d:
    del d['docProps/thumbnail.jpeg']
    from lxml import etree
    rr = etree.fromstring(d['_rels/.rels'])
    for r in list(rr):
        if 'thumbnail' in (r.get('Target', '') or ''):
            rr.remove(r)
    d['_rels/.rels'] = etree.tostring(rr, xml_declaration=True, encoding='UTF-8', standalone=True)

# Save
with _zf.ZipFile(OUTPUT, 'w') as z:
    for n, c in d.items():
        z.writestr(n, c)

os.remove(tmp)
print("✅ Готово: %s" % OUTPUT)
print("   Размер: %d байт" % os.path.getsize(OUTPUT))
