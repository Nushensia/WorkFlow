#!/usr/bin/env python3
"""
Создание DOCX-шаблона паспорта вырезки металла.
Обновлённая версия по результатам анализа исходного паспорта П-674.
— Титул: 2 строки 24pt
— Оглавление на 7 разделов
— Приложения с маркером по правому краю
— Характеристика — 14pt
— Таблица хим. состава — табуляцией (не Word-таблица), с элементом C
"""

import os, zipfile
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUTPUT = os.path.join(os.path.dirname(__file__), "Шаблон_паспорта.docx")

# ═══════════════════════════════════════════════
# 1. Документ
# ═══════════════════════════════════════════════
# Глобальный счётчик для нумерации полей TOC в lxml-части
_toc_counter = [0]

def make_field(para, fld_type, instr_text):
    """Добавить в параграф поле Word (begin/instrText/separate/end)."""
    r1 = OxmlElement('w:r')
    f1 = OxmlElement('w:fldChar')
    f1.set(qn('w:fldCharType'), 'begin')
    r1.append(f1)
    para._p.append(r1)

    r2 = OxmlElement('w:r')
    r2Pr = OxmlElement('w:rPr')
    r2.set(qn('w:rsidR'), '00F00001')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = instr_text
    r2.append(r2Pr)
    r2.append(instr)
    para._p.append(r2)

    r3 = OxmlElement('w:r')
    r3Pr = OxmlElement('w:rPr')
    f3 = OxmlElement('w:fldChar')
    f3.set(qn('w:fldCharType'), 'separate')
    r3.append(r3Pr)
    r3.append(f3)
    para._p.append(r3)

    r4 = OxmlElement('w:r')
    r4Pr = OxmlElement('w:rPr')
    t4 = OxmlElement('w:t')
    t4.text = '[ Обновить поле — правая кнопка мыши → Обновить поле ]'
    r4.append(r4Pr)
    r4.append(t4)
    para._p.append(r4)

    r5 = OxmlElement('w:r')
    r5Pr = OxmlElement('w:rPr')
    f5 = OxmlElement('w:fldChar')
    f5.set(qn('w:fldCharType'), 'end')
    r5.append(r5Pr)
    r5.append(f5)
    para._p.append(r5)
doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(1.5)

sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sty.paragraph_format.space_before = Pt(0)
sty.paragraph_format.space_after = Pt(0)


def p_center(text, bold=False, size=12, sb=0, sa=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0 if size <= 11 else 1.5
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p


def p_left(text, bold=False, size=12, sa=0, indent=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p


def p_right(text, bold=False, size=12, sa=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p


def p_just(text, bold=False, size=12, sa=0, italic=False, indent=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    return p


def mixed(label, placeholder, bold_label=True, size=12):
    """Поле с жирной меткой и обычным плейсхолдером."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    r1.bold = bold_label
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(size)
    r2 = p.add_run(placeholder)
    r2.bold = False
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(size)
    return p


def org_header():
    """Шапка организации — повторяется перед каждым протоколом."""
    p_center('Акционерное общество «НПО «Ленкор»', bold=True, size=12, sa=0)
    p_center('Лаборатория разрушающих методов контроля', bold=True, size=12, sa=0)
    p_center('Свидетельство об аккредитации', bold=False, size=12, sa=0)
    p_center('ОАО «НТЦ «Промышленная безопасность» № ИЛ/ЛРИ-02436', bold=False, size=12, sa=0)
    p_center('Адрес: Россия, 192236, г. Санкт-Петербург, ул. Белы Куна, д. 30, литера А, пом. 25-Н, офис 1408 тел. 8(812) 335-13-27; E-mail: office@npo-lencor.ru', bold=False, size=12, sa=12)


def heading1(text):
    """Заголовок стиля Heading 1 — для автоподбора TOC."""
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(16)
        r.bold = True
    return h


# ═══════════════════════════════════════════════
# ТИТУЛ — 2 строки 24pt
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p_center('ПАСПОРТ', bold=True, size=24, sa=0)
p_center('ВЫРЕЗКИ МЕТАЛЛА № [номер]', bold=True, size=24, sa=24)

# ═══════════════════════════════════════════════
# ОГЛАВЛЕНИЕ — настоящее поле TOC Word
# ═══════════════════════════════════════════════
p_left('Оглавление', bold=True, size=14, sa=6)

# Параграф с полем TOC
toc_para = doc.add_paragraph()
toc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
toc_para.paragraph_format.line_spacing = 1.5
make_field(toc_para, 'begin', 'TOC \\o "1-3" \\h \\z \\u')

doc.add_page_break()

# ═══════════════════════════════════════════════
# ХАРАКТЕРИСТИКА — 14pt
# ═══════════════════════════════════════════════
p_left('Краткая характеристика и эксплуатационные данные образца', bold=False, size=14, sa=6)

p_left('Название аппарата, трубопровода, детали: [наименование]', bold=False, size=14, sa=0)
p_left('Материал:', bold=False, size=14, sa=0)
p_left('а) по паспорту [марка по паспорту]', bold=False, size=14, sa=0)
p_left('б) фактический [марка фактическая]', bold=False, size=14, sa=0)
doc.add_paragraph()
p_left('1. Условия работы:', bold=False, size=14, sa=0)
p_left('а) среда: [среда]', bold=False, size=14, sa=0)
p_left('б) температура расчетная: [T_расч]°C', bold=False, size=14, sa=0)
p_left('   температура рабочая: [T_раб]°C', bold=False, size=14, sa=0)
p_left('в) давление расчетное: [P_расч]', bold=False, size=14, sa=0)
p_left('   давление рабочее: [P_раб]', bold=False, size=14, sa=0)
p_left('г) срок отработки металла: [срок] лет (с [год] года)', bold=False, size=14, sa=0)

p_left('2. Место вырезки образцов (эскиз вырезки образцов):', bold=False, size=14, sa=6)

p_center('[Место для вставки эскиза вырезки]', italic=True, size=12, sa=6)
p_center('Рисунок 1 — Схема вырезки образца', bold=False, size=12, sa=12)

p_left('3. Способ вырезки, маркировка и размеры вырезанного образца:', bold=False, size=14, sa=6)
p_just('[описание вырезки]', bold=False, size=12, sa=12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 1
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 1', bold=True, size=14, sa=6)
heading1('Фотографии образца')

p_center('[Место для вставки фотографий образца]', italic=True, size=12, sa=12)
p_center('Рисунок 1 — Фотографии вырезок металла П-[номер]', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 2 — ВИК
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 2', bold=True, size=14, sa=6)
heading1('Визуально-измерительный контроль (ВИК)')

org_header()
p_center('Протокол по визуально-измерительному контролю № [номер]/ВИК-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=12)

mixed('1.\tНаименование Заказчика испытаний: ', '[заказчик]')
mixed('2.\tОбъект контроля: ', '[объект]')
mixed('3.\tОснование для контроля: ', '[основание]')
mixed('4.\tЗадача контроля: ', '[задача]')
mixed('5.\tМетод контроля: ', '[метод со ссылками на НД]')
mixed('6.\tСредство контроля: ', '[средство]')
doc.add_paragraph()
p_just('7.\tРезультаты осмотра:', bold=True, sa=6)
p_just('а) [результаты осмотра]', sa=6, indent=1.0)
p_center('[Место для вставки фотографий ВИК]', italic=True, size=12, sa=6)
p_center('Рисунок 1 — Фотографии вырезок металла П-[номер]', bold=False, size=12, sa=12)

doc.add_paragraph()
p_left('Специалист ВИК', bold=False, size=12, sa=0)
p_left('Удостоверение о повышении квалификации № [номер удостоверения] от [дата]', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО специалиста]', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 3 — ХИМИЧЕСКИЙ АНАЛИЗ
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 3', bold=True, size=14, sa=6)
heading1('Химический анализ')

org_header()
p_center('Протокол химических испытаний № [номер]/Х-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=12)

mixed('1.\tНаименование Заказчика испытаний: ', '[заказчик]')
mixed('2.\tНаименование конкретной испытуемой продукции: ', '[продукция]')
mixed('3.\tКраткая характеристика испытуемого образца: ', '[образец]')
p_just('4.\tВид испытаний: количественный анализ химического состава материала образцов.', sa=0)
mixed('5.\tНормативные документы, использованные при испытаниях, в т. ч. методики: ', '[нормативные документы]')
mixed('6.\tИспытательное оборудование: ', '[оборудование]')
mixed('7.\tКоличество испытанных образцов и даты начала и окончания проведения испытаний: ', '[образцы и даты]')
doc.add_paragraph()

p_just('8.\tУсловия проведения испытаний:', bold=True, sa=6)
p_just('температура окружающего воздуха, °C    [t];', sa=0, indent=1.0)
p_just('атмосферное давление, мм рт. ст.       [P];', sa=0, indent=1.0)
p_just('относительная влажность воздуха, %     [влажность].', sa=12, indent=1.0)

p_just('9.\tРезультаты испытаний представлены в таблице:', bold=True, sa=6)

# Подпись таблицы — по левому краю 11pt
p_left('Таблица 1 — Результаты химического анализа образца и сравнение с требованиями на материал', bold=False, size=11, sa=6)

# Таблица хим. состава — табулированными параграфами (как в исходнике)
# Строка: элементы — 10pt полужирные
def chem_row(cells, bold=False):
    """Одна строка таблицы хим. состава — табулированные ячейки."""
    txt = '\t'.join(str(c) for c in cells)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(txt)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p

# Elements in original: Fe, C, Si, Mn, Cr, Mo, S, P, Ni, Cu (10 elements + Образец + Массовая доля %)
chem_headers = ['Образец', 'Массовая доля %', 'Fe', 'C', 'Si', 'Mn', 'Cr', 'Mo', 'S', 'P', 'Ni', 'Cu']
chem_data_1 = ['П-[номер]', 'осн.', '[C]', '[Si]', '[Mn]', '[Cr]', '[Mo]', '[S]', '[P]', '[Ni]', '[Cu]']
# Skip first two header cells for element row
chem_elements = ['', ''] + chem_headers[2:]

chem_row(chem_headers, bold=True)  # row 0: headers bold
# separator empty paragraph
doc.add_paragraph()
chem_row(chem_elements, bold=True)   # row 1: elements bold
chem_row(chem_data_1)                # row 2: data
# separator
doc.add_paragraph()
chem_row(['[стандарт 1]', 'осн.', '[C]', '[Si]', '[Mn]', '[Cr]', '[Mo]', '≤[S]', '≤[P]', '≤[Ni]', '≤[Cu]'])  # row 3
# separator
doc.add_paragraph()
chem_row(['[стандарт 2]', 'осн.', '[C]', '[Si]', '[Mn]', '[Cr]', '[Mo]', '≤[S]', '≤[P]', '≤[Ni]', '≤[Cu]'])  # row 4

doc.add_paragraph()
p_just('Заключение:', bold=True, sa=6)
p_just('По результатам химического анализа: [заключение]', sa=12)

p_left('Выполнил:', bold=False, size=12, sa=6)
doc.add_paragraph()
p_left('Руководитель лаборатории разрушающего контроля', bold=False, size=12, sa=0)
p_left('___________________________  [подпись]', bold=False, size=12, sa=0)
p_left('[ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 4 — МЕХАНИЧЕСКИЕ ИСПЫТАНИЯ
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 4', bold=True, size=14, sa=6)
heading1('Механические испытания')


def mech_table(headers, data_rows, size=9):
    """Универсальная таблица для механики."""
    nrows = 1 + len(data_rows)
    ncols = len(headers)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(t.cell(0, i), h, bold=True, size=size)
        shade_cell(t.cell(0, i))
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            set_cell_text(t.cell(ri + 1, ci), val, bold=False, size=size)
    return t


def set_cell_text(cell, text, bold=False, size=9, align='center'):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def shade_cell(cell, color='D9E2F3'):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def mech_protocol(num_label, title, test_type, size_headers, size_data, dim_rows, res_rows):
    """Один протокол внутри механики."""
    org_header()
    p_center(f'Протокол испытаний № [номер]/И/{num_label}-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=6)
    p_just(f'Вид испытаний: {test_type}', sa=6)
    p_just('Окружающая среда: [температура], относительная влажность [влажность]%', sa=6)
    p_just('Испытательное оборудование: [оборудование, № свидетельства, срок действия]', sa=12)
    p_just('Результаты испытаний:', bold=True, sa=6)

    p_center(f'Таблица {title}1 — Размеры образцов', bold=False, size=11, sa=6)
    dim_headers_mech = ['Маркировка\nобразца', 'Диаметр\nd, мм', 'Длина раб.\nчасти L, мм', 'Диаметр\nшейки dш, мм', 'Конечная длина\nраб. части, мм']
    mech_table(dim_headers_mech, dim_rows, size=size_data)

    doc.add_paragraph()
    p_center(f'Таблица {title}2 — Результаты испытаний на растяжение', bold=False, size=11, sa=6)
    res_headers_mech = ['Наименование\nобразца', 'Маркировка\nобразца', 'Температура\nиспытаний, °С', 'Предел\nтекучести\nσ₀.₂, МПа', 'Временное\nсопротивление\nσв, МПа', 'Отн.\nудлинение\nδ₅, %', 'Отн.\nсужение\nΨ, %']
    mech_table(res_headers_mech, res_rows, size=size_data)

    doc.add_paragraph()
    p_just('Фотографии образцов после испытаний:', bold=True, sa=6)
    p_center('[Место для вставки фотографий]', italic=True, size=12, sa=6)
    p_center(f'Рисунок {title[0]} — Фотографии образцов после испытаний', bold=False, size=12, sa=12)

    p_left('Инженер-металловед', bold=False, size=12, sa=0)
    p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
    p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=6)
    p_left('Руководитель лаборатории разрушающего контроля', bold=False, size=12, sa=0)
    p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
    p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=12)


# --- Протокол 1: Растяжение комн ---
dim_rows_1 = [['[маркировка]', '[d]', '[L]', '[dш]', '[Lк]']]
res_rows_1 = [
    ['[наим]', '[марка]', '[T]', '[σ₀.₂]', '[σв]', '[δ₅]', '[Ψ]'],
    ['[стандарт]', '[марка]', '[T]', '≥[σ₀.₂]', '≥[σв]', '≥[δ₅]', '≥[Ψ]'],
]
mech_protocol('1', '1', 'Испытания на статическое растяжение', 9, 9, dim_rows_1, res_rows_1)
doc.add_page_break()

# --- Протокол 2: Ударный изгиб ---
org_header()
p_center('Протокол испытаний № [номер]/И/2-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=6)
p_just('Вид испытаний: Испытания на ударный изгиб образцов с U-образным надрезом', sa=6)
p_just('Окружающая среда: [температура], относительная влажность [влажность]%', sa=6)
p_just('Испытательное оборудование: [оборудование]', sa=12)
p_just('Результаты испытаний:', bold=True, sa=6)

p_center('Таблица 3 — Результаты испытаний на ударный изгиб', bold=False, size=11, sa=6)

impact_headers = ['Шифр\nпартии', 'Маркировка\nобразца', 'Ширина\nсечения\nB, мм', 'Высота\nсечения\nH, мм', 'Высота раб.\nсечения\nH₁, мм', 'Энергия\nудара,\nДж', 'Средняя\nэнергия\nудара, Дж', 'Ударная\nвязкость\nKCU,\nДж/см²', 'Средняя\nударная\nвязкость\nKCU,\nДж/см²']
impact_rows = [
    ['[шифр]', '[марка]', '[B]', '[H]', '[H₁]', '[энергия]', '[ср. эн.]', '[KCU]', '[ср. KCU]'],
    ['[стандарт]', '', '≥[B]', '', '', '', '', '≥[KCU]', ''],
]
mech_table(impact_headers, impact_rows, size=8)

doc.add_paragraph()
p_just('Фотографии образцов после испытаний:', bold=True, sa=6)
p_center('[Место для вставки фотографий]', italic=True, size=12, sa=12)

p_left('Инженер-металловед', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=6)
p_left('Руководитель лаборатории разрушающего контроля', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=12)
doc.add_page_break()

# --- Протокол 3: Растяжение Tраб ---
dim_rows_3 = [['[маркировка]', '[d]', '[L]', '[dш]', '[Lк]']]
res_rows_3 = [
    ['[наим]', '[марка]', '[T]', '[σ₀.₂]', '[σв]', '[δ₅]', '[Ψ]'],
    ['[стандарт]', '[марка]', '[T]', '≥[σ₀.₂]', '≥[σв] / [диапазон]', '≥[δ₅]', '≥[Ψ]'],
]
mech_protocol('3', '4', 'Испытания на статическое растяжение при [T]°С', 9, 9, dim_rows_3, res_rows_3)

p_center('Приложение 1. Диаграммы деформирования ([N] шт.).', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 5 — МЕТАЛЛОГРАФИЯ
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 5', bold=True, size=14, sa=6)
heading1('Металлография')

org_header()
p_center('ПРОТОКОЛ МЕТАЛЛОГРАФИЧЕСКОГО ИССЛЕДОВАНИЯ', bold=True, size=12, sa=6)
p_center('№ [номер]/М-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=12)

mixed('1.\tНаименование Заказчика испытаний: ', '[заказчик]')
mixed('2.\tВид испытаний: ', 'Металлографический анализ структуры на плоских сечениях.')
mixed('3.\tЦель испытаний: ', '[цель]')
mixed('4.\tНормативная документация, использованная при испытаниях: ', '[нормативные документы]')
doc.add_paragraph()

p_just('5.\tРезультаты испытаний:', bold=True, sa=6)
p_center('Рисунок 1 — Микроструктура образца П-[номер]-[N], шлиф не протравлен, х100', bold=False, size=10, sa=6)
p_center('[Место для вставки микрофотографий]', italic=True, size=12, sa=6)
p_center('а) х100     б) х200     в) х500     г) х1000', bold=False, size=10, sa=12)
doc.add_paragraph()
p_just('Выводы:', bold=True, sa=12)

# Таблица микротвердости
p_center('Таблица 1 — Результаты измерения микротвердости образца П-[номер]', bold=False, size=10, sa=6)

micro_h = ['Участок измерений', 'Интервал значений\nHV 0.2', 'Среднее\nHV 0.2', 'HB\n(перевод по ASTM A370)']
micro_d = [
    ['Наружный край', '[HV интервал]', '[средняя HV]', '[HB]'],
    ['Центральная часть', '[HV интервал]', '[средняя HV]', '[HB]'],
    ['Внутренний край', '[HV интервал]', '[средняя HV]', '[HB]'],
]
tbl_hv = doc.add_table(rows=4, cols=4)
tbl_hv.style = 'Table Grid'
tbl_hv.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(micro_h):
    set_cell_text(tbl_hv.cell(0, i), h, bold=True, size=9)
    shade_cell(tbl_hv.cell(0, i))
for ri, row in enumerate(micro_d):
    for ci, val in enumerate(row):
        set_cell_text(tbl_hv.cell(ri + 1, ci), val, bold=False, size=9)
for row in tbl_hv.rows:
    for idx, w in enumerate([Cm(3.0), Cm(3.0), Cm(2.5), Cm(3.5)]):
        row.cells[idx].width = w

doc.add_paragraph()
p_just('Результаты РСМА анализа:', bold=True, sa=6)

# Таблица РСМА 1
p_center('Таблица 2 — Результаты РСМА анализа.', bold=False, size=10, sa=6)
rsma_h = ['№ спектра', 'Fe', 'Mn', 'Si', 'Cr', 'Mo', 'Al']
rsma_d = [
    ['[спектр 1]', 'ост.', '[Mn]', '[Si]', '[Cr]', '[Mo]', '[Al]'],
    ['[спектр 2]', 'ост.', '[Mn]', '[Si]', '[Cr]', '[Mo]', '[Al]'],
]
mech_table(rsma_h, rsma_d, size=9)

# Таблица РСМА 2
doc.add_paragraph()
p_center('Таблица 3 — Результаты РСМА анализа.', bold=False, size=10, sa=6)
mech_table(rsma_h, rsma_d, size=9)

doc.add_paragraph()
p_left('Инженер-металловед', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=6)
p_left('Специалист по металлографическим исследованиям', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 6 — ТВЁРДОСТЬ
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 6', bold=True, size=14, sa=6)
heading1('Твердость')

org_header()
p_center('Протокол измерений твердости металла № [номер]/Т/[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=6)

mixed('Объект исследования: ', '[объект]')
mixed('Вид испытаний: ', 'Измерение твердости образцов металла.')
mixed('Окружающая среда: ', '[температура], относительная влажность [влажность]%')
mixed('Нормативные документы, использованные при испытаниях, в т. ч. методики: ', '[нормативные документы]')
mixed('Испытательное оборудование: ', '[оборудование, № свидетельства, срок действия]')
doc.add_paragraph()

# Таблица твёрдости — трёхзонная
p_center('Таблица — Результаты измерения твердости металла', bold=False, size=11, sa=6)

# 3 зоны × (Номер точки | HRB | HB) = 9 columns
tbl_hard = doc.add_table(rows=15, cols=9)
tbl_hard.style = 'Table Grid'
tbl_hard.alignment = WD_TABLE_ALIGNMENT.CENTER

# Row 0: zone headers (merged)
zone_labels = ['Наружный край', 'Центральная часть', 'Внутренний край']
zone_ranges = [(0, 2), (3, 5), (6, 8)]
for zi, (label, (cs, ce)) in enumerate(zip(zone_labels, zone_ranges)):
    for c in range(cs, ce + 1):
        set_cell_text(tbl_hard.cell(0, c), label, bold=True, size=9)
        shade_cell(tbl_hard.cell(0, c))
    if cs < ce:
        tbl_hard.cell(0, cs).merge(tbl_hard.cell(0, ce))

# Row 1: sub-headers
sub = ['№ точки', 'HRB', 'HB'] * 3
for c, h in enumerate(sub):
    set_cell_text(tbl_hard.cell(1, c), h, bold=True, size=9)
    shade_cell(tbl_hard.cell(1, c))

# Rows 2-7: data (6 замеров)
for r in range(2, 8):
    # Each zone has № точки, HRB, HB
    for zi in range(3):
        base = zi * 3
        set_cell_text(tbl_hard.cell(r, base), str(r - 1), bold=False, size=9)
        set_cell_text(tbl_hard.cell(r, base + 1), '[HRB]', bold=False, size=9)
        set_cell_text(tbl_hard.cell(r, base + 2), '[HB]', bold=False, size=9)

# Rows 8-10: Среднее, Min, Max
summary_labels = ['Среднее', 'Min', 'Max']
for si, slabel in enumerate(summary_labels):
    r = 8 + si
    for zi in range(3):
        base = zi * 3
        set_cell_text(tbl_hard.cell(r, base), slabel, bold=True, size=9)
        set_cell_text(tbl_hard.cell(r, base + 1), '[HRB]', bold=True, size=9)
        set_cell_text(tbl_hard.cell(r, base + 2), '[HB]', bold=True, size=9)

# Ширина колонок
for row in tbl_hard.rows:
    for idx, w in enumerate([Cm(1.2), Cm(1.2), Cm(1.2)] * 3):
        row.cells[idx].width = w

doc.add_paragraph()
p_just('10.2. Фотография образцов измерения твердости', bold=True, sa=6)
p_center('[Место для вставки фотографий]', italic=True, size=12, sa=12)

p_just('Заключение по результатам испытаний: [заключение]', sa=6)
p_just('Настоящий протокол распространяется только на образцы, подвергнутые испытанию.', sa=12)

p_left('Руководитель лаборатории разрушающего контроля', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=6)
p_left('Инженер-металловед', bold=False, size=12, sa=0)
p_left('___________________________  [ФИО]', bold=False, size=12, sa=0)
p_left('Квалификационное удостоверение № [номер] до [дата]', bold=False, size=12, sa=12)
doc.add_page_break()

# ═══════════════════════════════════════════════
# ПРИЛОЖЕНИЕ 7 — ФАЗОВЫЙ АНАЛИЗ ОСАДКА (чистый python-docx, без lxml-постпроцессинга)
# ═══════════════════════════════════════════════
p_right('ПРИЛОЖЕНИЕ 7', bold=True, size=14, sa=6)
heading1('Фазовый анализ осадка')

# --- Создаём таблицу-рамку 2×1 python-docx ---
fa_border = doc.add_table(rows=2, cols=1)
fa_border.alignment = WD_TABLE_ALIGNMENT.CENTER
# Устанавливаем границы таблицы через lxml (единственное место)
tbl_pr = fa_border._tbl.find(qn('w:tblPr'))
if tbl_pr is None:
    tbl_pr = OxmlElement('w:tblPr')
    fa_border._tbl.insert(0, tbl_pr)
tbl_borders = OxmlElement('w:tblBorders')
for side in ['top', 'left', 'bottom', 'right']:
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '18')
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), 'auto')
    tbl_borders.append(b)
tbl_pr.append(tbl_borders)
# Ширина таблицы на всю страницу
tbl_w = OxmlElement('w:tblW')
tbl_w.set(qn('w:w'), '0')
tbl_w.set(qn('w:type'), 'auto')
tbl_pr.append(tbl_w)

# --- Строка 0: Шапка организации ---
c0 = fa_border.cell(0, 0)
# Очищаем параграф по умолчанию
c0.paragraphs[0].clear()
c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# Добавляем строки шапки
for txt, bold in [('Акционерное общество «НПО «Ленкор»', True),
                   ('Лаборатория разрушающих методов контроля', False),
                   ('Свидетельство об аккредитации', False),
                   ('ОАО «НТЦ «Промышленная безопасность» № ИЛ/ЛРИ-02436', False)]:
    p = c0.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(txt)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
# Удаляем пустой первый параграф
c0.paragraphs[0]._element.getparent().remove(c0.paragraphs[0]._element)

# --- Строка 1: Контент протокола ФА ---
c1 = fa_border.cell(1, 0)
# Очищаем параграф по умолчанию
c1.paragraphs[0].clear()

def cell_center(cell, text, bold=False, size=12, sb=0, sa=0, italic=False):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0 if size <= 11 else 1.5
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    return p

def cell_left(cell, text, bold=False, size=12, sa=0):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    return p

def cell_just(cell, text, bold=False, size=12, sa=0, italic=False, indent=None):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    return p

def cell_mixed(cell, label, placeholder, size=12):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label); r1.bold = True
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(size)
    r2 = p.add_run(placeholder); r2.bold = False
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(size)
    return p

# Контент
cell_center(c1, 'ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА', bold=True, size=14, sa=6)
cell_left(c1, '№ [номер]/КА-[год] от «[дд]» [месяц] [гггг] г.', bold=True, size=12, sa=12)
cell_mixed(c1, '1.\tНаименование Заказчика испытаний: ', '[заказчик]')
cell_mixed(c1, '2.\tНаименование конкретной испытуемой продукции: ', '[продукция]')
cell_mixed(c1, '3.\tКраткая характеристика испытуемого образца: ', '[марка стали]')
cell_just(c1, '4.\tВид испытаний: Физико-химический фазовый анализ.', sa=0)
cell_mixed(c1, '5.\tНормативные документы, использованные при испытаниях, в т. ч. методики: ', '[нормативные документы]')
cell_mixed(c1, '6.\tИспытательное оборудование: ', '[оборудование]')
cell_mixed(c1, '7.\tКоличество испытанных образцов и даты начала и окончания проведения испытаний: ', '[образцы и даты]')
c1.add_paragraph()
cell_just(c1, '8.\tРезультаты испытаний:', bold=True, sa=6)
cell_just(c1, '[условия_электрохимического_травления]', sa=6, italic=True)
cell_just(c1, '[режимы_рентгенографирования]', sa=12, italic=True)
cell_just(c1, '[описание_фазового_состава]', sa=6, italic=True)
cell_center(c1, 'Таблица 1 — Фазовый состав образца П-[номер]', bold=False, size=11, sa=6)

# Таблица фаз — добавляем В ячейку c1
tbl_phase = c1.add_table(rows=3, cols=7)
tbl_phase.style = 'Table Grid'
tbl_phase.alignment = WD_TABLE_ALIGNMENT.CENTER
phase_h = ['№', 'Phase name', 'Chemical\nformula', 'Crystal\nsystem', 'Space\ngroup', 'Lattice\nparameters, Å', 'DB card\nnumber']
for i, h in enumerate(phase_h):
    set_cell_text(tbl_phase.cell(0, i), h, bold=True, size=9)
    shade_cell(tbl_phase.cell(0, i))
for i, val in enumerate(['1', '[Phase name]', '[Formula]', '[Sys]', '[SG]', '[a,b,c,α,β,γ]', '[DB#]']):
    align = 'center' if i != 1 else 'left'
    set_cell_text(tbl_phase.cell(1, i), val, bold=False, size=9, align=align)
for i in range(7):
    set_cell_text(tbl_phase.cell(2, i), '...', bold=False, size=9)
for row in tbl_phase.rows:
    for idx, w in enumerate([Cm(0.8), Cm(3.0), Cm(2.5), Cm(2.0), Cm(2.5), Cm(3.5), Cm(2.5)]):
        row.cells[idx].width = w

c1.add_paragraph()
cell_center(c1, '[Место для вставки рисунка — Дифрактограмма]', italic=True, size=12, sa=6)
cell_center(c1, 'Рисунок 1 — Дифрактограмма образца П-[номер]', bold=False, size=12, sa=12)
cell_just(c1, 'Заключение:', bold=True, sa=6)
cell_just(c1, 'По результатам фазового анализа состав образца представлен следующими фазами:', sa=6, indent=1.0)
cell_just(c1, '• [карбиды типа M₂₃C₆ и M₆C]', sa=0, indent=1.5)
cell_just(c1, '• [оксиды типа M₃O₄]', sa=0, indent=1.5)
cell_just(c1, '• [прочие фазы]', sa=6, indent=1.5)
p_note = c1.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_note.paragraph_format.line_spacing = 1.5
p_note.paragraph_format.space_before = Pt(0)
p_note.paragraph_format.space_after = Pt(0)
rn = p_note.add_run('Примечание:')
rn.bold = True; rn.italic = True
rn.font.name = 'Times New Roman'; rn.font.size = Pt(12)
cell_just(c1, '[текст примечания]', sa=12, italic=True)
c1.add_paragraph()
cell_left(c1, 'Заместитель технического директора по развитию и науке', bold=False, size=12, sa=0)
c1.add_paragraph()
cell_left(c1, '___________________________  [подпись]', bold=False, size=12, sa=0)
cell_left(c1, '[ФИО]', bold=False, size=12, sa=12)

doc.add_paragraph()  # отступ после таблицы-рамки

# ═══════════════════════════════════════════════
# Сохраняем через python-docx
# ═══════════════════════════════════════════════
TMP = OUTPUT.replace('.docx', '_tmp.docx')
doc.save(TMP)

# ═══════════════════════════════════════════════
# 2. Footer с PAGE/NUMPAGES через lxml
# ═══════════════════════════════════════════════
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_r = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
ns_rel = 'http://schemas.openxmlformats.org/package/2006/relationships'
ns_ct = 'http://schemas.openxmlformats.org/package/2006/content-types'
W = f'{{{ns_w}}}'

with zipfile.ZipFile(TMP, 'r') as z:
    data = {n: z.read(n) for n in z.namelist()}

doc_xml = etree.fromstring(data['word/document.xml'])
body = doc_xml.find(f'{W}body')
sectPr = body.find(f'{W}sectPr')

# ═══════════════════════════════════════════════
# 2b. Footer с PAGE/NUMPAGES
footer_xml = etree.Element(f'{W}ftr')
fp = etree.SubElement(footer_xml, f'{W}p')
fpp = etree.SubElement(fp, f'{W}pPr')
fjc = etree.SubElement(fpp, f'{W}jc')
fjc.set(f'{W}val', 'center')
fsp = etree.SubElement(fpp, f'{W}spacing')
fsp.set(f'{W}line', '240')
fsp.set(f'{W}lineRule', 'auto')


def add_field(para, field_name):
    """Добавить поле PAGE или NUMPAGES."""
    r1 = etree.SubElement(para, f'{W}r')
    rp1 = etree.SubElement(r1, f'{W}rPr')
    etree.SubElement(rp1, f'{W}b')
    etree.SubElement(rp1, f'{W}bCs')
    rf = etree.SubElement(rp1, f'{W}rFonts')
    rf.set(f'{W}ascii', 'Times New Roman')
    rf.set(f'{W}hAnsi', 'Times New Roman')
    rs = etree.SubElement(rp1, f'{W}sz')
    rs.set(f'{W}val', '20')
    etree.SubElement(r1, f'{W}fldChar').set(f'{W}fldCharType', 'begin')

    r2 = etree.SubElement(para, f'{W}r')
    rp2 = etree.SubElement(r2, f'{W}rPr')
    etree.SubElement(rp2, f'{W}b')
    etree.SubElement(rp2, f'{W}bCs')
    rf2 = etree.SubElement(rp2, f'{W}rFonts')
    rf2.set(f'{W}ascii', 'Times New Roman')
    rf2.set(f'{W}hAnsi', 'Times New Roman')
    rs2 = etree.SubElement(rp2, f'{W}sz')
    rs2.set(f'{W}val', '20')
    instr = etree.SubElement(r2, f'{W}instrText')
    instr.text = field_name

    r3 = etree.SubElement(para, f'{W}r')
    rp3 = etree.SubElement(r3, f'{W}rPr')
    etree.SubElement(rp3, f'{W}b')
    etree.SubElement(rp3, f'{W}bCs')
    rf3 = etree.SubElement(rp3, f'{W}rFonts')
    rf3.set(f'{W}ascii', 'Times New Roman')
    rf3.set(f'{W}hAnsi', 'Times New Roman')
    rs3 = etree.SubElement(rp3, f'{W}sz')
    rs3.set(f'{W}val', '20')
    etree.SubElement(r3, f'{W}fldChar').set(f'{W}fldCharType', 'separate')

    r4 = etree.SubElement(para, f'{W}r')
    rp4 = etree.SubElement(r4, f'{W}rPr')
    etree.SubElement(rp4, f'{W}b')
    etree.SubElement(rp4, f'{W}bCs')
    rf4 = etree.SubElement(rp4, f'{W}rFonts')
    rf4.set(f'{W}ascii', 'Times New Roman')
    rf4.set(f'{W}hAnsi', 'Times New Roman')
    rs4 = etree.SubElement(rp4, f'{W}sz')
    rs4.set(f'{W}val', '20')
    tv = etree.SubElement(r4, f'{W}t')
    tv.text = '1'

    r5 = etree.SubElement(para, f'{W}r')
    rp5 = etree.SubElement(r5, f'{W}rPr')
    etree.SubElement(rp5, f'{W}b')
    etree.SubElement(rp5, f'{W}bCs')
    rf5 = etree.SubElement(rp5, f'{W}rFonts')
    rf5.set(f'{W}ascii', 'Times New Roman')
    rf5.set(f'{W}hAnsi', 'Times New Roman')
    rs5 = etree.SubElement(rp5, f'{W}sz')
    rs5.set(f'{W}val', '20')
    etree.SubElement(r5, f'{W}fldChar').set(f'{W}fldCharType', 'end')


# Собираем footer: "Лист " + PAGE + " из " + NUMPAGES
parts = [
    ('Лист ', False),
    ('PAGE', True),
    (' из ', False),
    ('NUMPAGES', True),
]
for text, is_field in parts:
    if is_field:
        add_field(fp, text)
    else:
        r = etree.SubElement(fp, f'{W}r')
        rp = etree.SubElement(r, f'{W}rPr')
        rf = etree.SubElement(rp, f'{W}rFonts')
        rf.set(f'{W}ascii', 'Times New Roman')
        rf.set(f'{W}hAnsi', 'Times New Roman')
        rs = etree.SubElement(rp, f'{W}sz')
        rs.set(f'{W}val', '20')
        t = etree.SubElement(r, f'{W}t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text

data['word/footer1.xml'] = etree.tostring(footer_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

# Rels
rels = etree.fromstring(data['word/_rels/document.xml.rels'])
footer_rel_exists = any('footer' in (r.get('Target', '') or '') for r in rels)
if not footer_rel_exists:
    existing_ids = [int(r.get('Id', 'rId0')[3:]) for r in rels]
    max_id = max(existing_ids) if existing_ids else 0
    new_rel = etree.SubElement(rels, f'{{{ns_rel}}}Relationship')
    new_rel.set('Id', f'rId{max_id + 1}')
    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer')
    new_rel.set('Target', 'footer1.xml')
    footer_rid = f'rId{max_id + 1}'
else:
    footer_rid = None
    for r in rels:
        if 'footer' in (r.get('Target', '') or ''):
            footer_rid = r.get('Id')
            break

if sectPr is not None and footer_rid:
    fr = sectPr.find(f'{W}footerReference')
    if fr is None:
        fr_new = etree.SubElement(sectPr, f'{W}footerReference')
        fr_new.set(f'{W}type', 'default')
        fr_new.set(f'{{{ns_r}}}id', footer_rid)
        pgSz = sectPr.find(f'{W}pgSz')
        if pgSz is not None:
            sectPr.remove(fr_new)
            pgSz.addnext(fr_new)

data['word/document.xml'] = etree.tostring(doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
data['word/_rels/document.xml.rels'] = etree.tostring(rels, xml_declaration=True, encoding='UTF-8', standalone=True)

# Content_Types
ct = etree.fromstring(data['[Content_Types].xml'])
if ct.find(f'{{{ns_ct}}}Override[@PartName="/word/footer1.xml"]') is None:
    ov = etree.SubElement(ct, f'{{{ns_ct}}}Override')
    ov.set('PartName', '/word/footer1.xml')
    ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml')
data['[Content_Types].xml'] = etree.tostring(ct, xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════
# 3. Очистка: удаляем stylesWithEffects, thumbnail, правим Content_Types и rels
# ═══════════════════════════════════════════════
for key in list(data.keys()):
    if 'stylesWithEffects' in key or 'thumbnail' in key.lower():
        del data[key]

# Удаляем из Content_Types ссылки на удалённые файлы
ct = etree.fromstring(data['[Content_Types].xml'])
for override in list(ct):
    pn = override.get('PartName', '')
    if 'stylesWithEffects' in pn or 'thumbnail' in pn.lower():
        ct.remove(override)
for default in list(ct):
    ext = default.get('Extension', '')
    if ext == 'jpeg' and 'thumbnail' in etree.tostring(ct, encoding='unicode'):
        # Удаляем default для jpeg если используется только для thumbnail
        all_uses = [ov for ov in ct if ov.get('PartName', '').endswith('.jpeg')]
        if not all_uses:
            ct.remove(default)
data['[Content_Types].xml'] = etree.tostring(ct, xml_declaration=True, encoding='UTF-8', standalone=True)

# Удаляем thumbnail из _rels/.rels
rels_root = etree.fromstring(data['_rels/.rels'])
for rel in list(rels_root):
    target = rel.get('Target', '')
    if 'thumbnail' in target.lower():
        rels_root.remove(rel)
data['_rels/.rels'] = etree.tostring(rels_root, xml_declaration=True, encoding='UTF-8', standalone=True)

# Сохраняем
tmp2 = TMP.replace('_tmp', '_tmp2')
with zipfile.ZipFile(tmp2, 'w') as z:
    for name, content in data.items():
        z.writestr(name, content)

os.replace(tmp2, OUTPUT)
os.remove(TMP)

print(f"✅ Шаблон паспорта сохранён: {OUTPUT}")
