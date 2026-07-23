import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, Twips, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.section import WD_ORIENT
import os


def make_border_xml(sz="18", color="auto", val="single"):
    """Create borders XML element with specified size (in eighths of pt)."""
    return (
        '<w:tcBorders %s>'
        '<w:top w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:left w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:bottom w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '<w:right w:val="%s" w:sz="%s" w:space="0" w:color="%s"/>'
        '</w:tcBorders>'
    ) % (nsdecls('w'), val, sz, color, val, sz, color, val, sz, color, val, sz, color)


def set_cell_border(cell, sz="18", color="auto", val="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(make_border_xml(sz, color, val))
    tcPr.append(tcBorders)


def set_cell_noborder(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders %s>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>' % nsdecls('w'))
    tcPr.append(tcBorders)


def set_cell_vertical_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)


def hide_cell_mark(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    hide = OxmlElement('w:hideMark')
    tcPr.append(hide)


def add_run(p, text, bold=False, italic=False, size=12, font_name='Times New Roman', color=None):
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def new_para(doc_or_cell, text='', bold=False, italic=False, size=12, alignment=None,
             space_before=0, space_after=0, left_indent=None, font_name='Times New Roman',
             first_line_indent=None):
    p = doc_or_cell.add_paragraph()
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, font_name=font_name)
    if alignment is not None:
        p.alignment = alignment
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_field(p, field_code, bold=False, font_name='Times New Roman', size=11):
    run1 = p.add_run()
    fldChar1 = parse_xml('<w:fldChar ' + nsdecls('w') + ' w:fldCharType="begin"/>')
    run1._r.append(fldChar1)

    run2 = p.add_run()
    instrText = parse_xml('<w:instrText ' + nsdecls('w') + ' xml:space="preserve"> ' + field_code + ' </w:instrText>')
    run2._r.append(instrText)

    run3 = p.add_run()
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(size * 2))
    rPr.append(szCs)
    run3._r.append(rPr)

    fldChar2 = parse_xml('<w:fldChar ' + nsdecls('w') + ' w:fldCharType="end"/>')
    run3._r.append(fldChar2)


# ============================================================
# CREATE DOCUMENT
# ============================================================
doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(3.0)
section.right_margin = Cm(1.5)
section.header_distance = Cm(1.4)
section.footer_distance = Cm(1.4)

# Default styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15

# ============================================================
# LOGO (Russian text only, from PDF)
# ============================================================
logo_path = r'D:\ИИ\Git Projetcs\Nushensia\WorkFlow\Паспорта\Исходные протоколы\logo_top.png'
print('Logo prepared (top region)')

# ============================================================
# OUTER TABLE (рамка: 2 строки, insideH = жирная черта)
# ============================================================
outer_table = doc.add_table(rows=2, cols=1)
outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tblPr = outer_table._tbl.tblPr
tblBorders = parse_xml(
    '<w:tblBorders %s>'
    '<w:top w:val="single" w:sz="18" w:space="0" w:color="auto"/>'
    '<w:left w:val="single" w:sz="18" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="single" w:sz="18" w:space="0" w:color="auto"/>'
    '<w:right w:val="single" w:sz="18" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="single" w:sz="18" w:space="0" w:color="auto"/>'
    '</w:tblBorders>' % nsdecls('w'))
tblPr.append(tblBorders)

tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'), '0')
tblW.set(qn('w:type'), 'auto')
tblPr.append(tblW)

# ============================================================
# ROW 1: ШАПКА (информация о компании по центру)
# ============================================================
header_cell = outer_table.cell(0, 0)
tcPr = header_cell._tc.get_or_add_tcPr()
tcBorders = parse_xml(make_border_xml("18", "auto", "single"))
tcPr.append(tcBorders)

# --- Nested header table (no borders, logo + company name) ---
hdr_table = header_cell.add_table(rows=1, cols=2)
hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tblPr_h = hdr_table._tbl.tblPr
tblBorders_h = parse_xml(
    '<w:tblBorders %s>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>' % nsdecls('w'))
tblPr_h.append(tblBorders_h)

# Grid columns
tblGrid_h = OxmlElement('w:tblGrid')
for w_val in [2466, 6519]:
    col = OxmlElement('w:gridCol')
    col.set(qn('w:w'), str(w_val))
    tblGrid_h.append(col)
hdr_table._tbl.insert(0, tblGrid_h)

# Cell 0: Logo (high-res PNG via add_picture)
hc0 = hdr_table.cell(0, 0)
set_cell_noborder(hc0)
set_cell_vertical_center(hc0)
p_logo = hc0.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

from docx.shared import Emu
run = p_logo.add_run()
# Only specify width, height auto-calculates to maintain aspect ratio
run.add_picture(logo_path, width=Emu(1409700))
print('Logo added (PNG, aspect ratio preserved)')

# Cell 1: Empty
hc1 = hdr_table.cell(0, 1)
set_cell_noborder(hc1)
set_cell_vertical_center(hc1)

# --- Centered paragraphs below the nested table ---
p = new_para(header_cell, 'Акционерное общество «НПО «Ленкор»',
             bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=2)

p = new_para(header_cell, 'Лаборатория разрушающих методов контроля',
             bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

p = new_para(header_cell, 'Свидетельство об аккредитации',
             size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

p = new_para(header_cell, 'ОАО «НТЦ «Промышленная безопасность» № ИЛ/ЛРИ-02436',
             size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

p = new_para(header_cell, 'Россия, 192236, г. Санкт-Петербург, ул. Белы Куна, д. 30, литера А, пом. 25-Н, офис 1408 тел. 8(812) 335-13-27; E-mail: office@npo-lencor.ru',
             size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# ROW 2: ОСНОВНОЕ СОДЕРЖИМОЕ
# ============================================================
outer_cell = outer_table.cell(1, 0)
tcPr2 = outer_cell._tc.get_or_add_tcPr()
tcBorders2 = parse_xml(make_border_xml("18", "auto", "single"))
tcPr2.append(tcBorders2)

# --- Title ---
p = new_para(outer_cell, 'ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА',
             bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=18, space_after=6)

# --- Protocol Number ---
p = new_para(outer_cell, '', size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
# Use multiple runs to handle special characters
run1 = p.add_run('№ ')
run1.font.size = Pt(14)
run1.font.name = 'Times New Roman'
run2 = p.add_run('{{НОМЕР_ПРОТОКОЛА}}')
run2.font.size = Pt(14)
run2.font.name = 'Times New Roman'
run3 = p.add_run('/КА-')
run3.font.size = Pt(14)
run3.font.name = 'Times New Roman'
run4 = p.add_run('{{ГОД}}')
run4.font.size = Pt(14)
run4.font.name = 'Times New Roman'
run5 = p.add_run(' от «')
run5.font.size = Pt(14)
run5.font.name = 'Times New Roman'
run6 = p.add_run('{{ДЕНЬ}}')
run6.font.size = Pt(14)
run6.font.name = 'Times New Roman'
run7 = p.add_run('» ')
run7.font.size = Pt(14)
run7.font.name = 'Times New Roman'
run8 = p.add_run('{{МЕСЯЦ_ПРОПИСЬЮ}}')
run8.font.size = Pt(14)
run8.font.name = 'Times New Roman'
run9 = p.add_run(' ')
run9.font.size = Pt(14)
run9.font.name = 'Times New Roman'
run10 = p.add_run('{{ГОД}}')
run10.font.size = Pt(14)
run10.font.name = 'Times New Roman'
run11 = p.add_run(' г.')
run11.font.size = Pt(14)
run11.font.name = 'Times New Roman'

# --- Sections with numbering ---
def add_num_item(cell, num, title, value='', indent_twips=709, hanging_twips=425):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Twips(indent_twips)
    p.paragraph_format.first_line_indent = Twips(-hanging_twips)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # Number prefix
    r = p.add_run(num + ' ')
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Title
    r = p.add_run(title + ' ')
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    # Value
    if value:
        r = p.add_run(value)
        r.font.size = Pt(12)
        r.font.name = 'Times New Roman'
    return p


def add_bullet_item(cell, text, indent_left=720, indent_hanging=83, size=12):
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Twips(indent_left)
    p.paragraph_format.first_line_indent = Twips(-indent_hanging)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run('\u2022 ' + text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p


add_num_item(outer_cell, '\u043f.1', 'Наименование Заказчика испытаний:', '{{ЗАКАЗЧИК}}')
add_num_item(outer_cell, '\u043f.2', 'Наименование конкретной испытуемой продукции:', '{{ОБРАЗЕЦ_ОПИСАНИЕ}}')
add_num_item(outer_cell, '\u043f.3', 'Краткая характеристика испытуемого образца:', '{{МАРКА_СТАЛИ}}')
add_num_item(outer_cell, '\u043f.4', 'Вид испытаний:', 'Физико-химический фазовый анализ.')

add_num_item(outer_cell, '\u043f.5', 'Нормативные документы, использованные при испытаниях, в т. ч. методики:')
add_bullet_item(outer_cell, 'Методика анализа фазового состава конструкционных наноматериалов методом рентгеновской дифрактометрии, свидетельство об аттестации \u211601.00225/206-03-2011 от 20.05.2011\u0433, регистрационный код ФР.1.31.2011.10209.')
add_bullet_item(outer_cell, 'Базы порошковых дифракционных стандартов COD (Crystallography Open Database), PDF-2 (Powder Diffraction File).')

add_num_item(outer_cell, '\u043f.6', 'Испытательное оборудование:')
add_bullet_item(outer_cell, 'Многофункциональный рентгеновский дифрактометр Rigaku Ultima IV с комплексом управляющих программ и обрабатывающим комплексом PDXL (X-ray Powder Diffraction Software). Свидетельство о поверке \u2116 {{НОМЕР_ПОВЕРКИ}}.')

add_num_item(outer_cell, '\u043f.7', 'Количество испытанных образцов и даты начала и окончания проведения испытаний:',
             '{{КОЛ_ВО_ОБРАЗЦОВ}}, дата испытаний {{ДАТА_НАЧАЛА}}\u2013{{ДАТА_ОКОНЧАНИЯ}}')

add_num_item(outer_cell, '\u043f.8', 'Результаты испытаний:')
add_bullet_item(outer_cell, 'Условия электрохимического травления: электролит {{ЭЛЕКТРОЛИТ}}, 1:20, \u201310 \u00b0C, плотность тока 2 \u0410/\u0434\u043c\u00b2, время проведения \u2014 2 \u0447;', indent_left=720)
add_bullet_item(outer_cell, '{{ОПИСАНИЕ_РЕНТГЕНОГРАФИРОВАНИЯ}}', indent_left=720)
add_bullet_item(outer_cell, 'Фазовый состав образца представлен следующими основными фазами: {{ОПИСАНИЕ_ФАЗ}}', indent_left=720)

# --- Phase Composition Table ---
p = new_para(outer_cell, 'Таблица 1 \u2014 Фазовый состав образца {{ОБОЗНАЧЕНИЕ}}',
             bold=True, size=11, space_before=12, space_after=6)
# Keep with next
pPr = p._p.get_or_add_pPr()
keepNext = OxmlElement('w:keepNext')
pPr.append(keepNext)

phase_table = outer_cell.add_table(rows=2, cols=7)
phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER
phase_table.style = 'Table Grid'

# Set table width to 100%
tblPr3 = phase_table._tbl.tblPr
tblW3 = OxmlElement('w:tblW')
tblW3.set(qn('w:w'), '5000')
tblW3.set(qn('w:type'), 'pct')
tblPr3.append(tblW3)

# Set column widths
tblGrid3 = OxmlElement('w:tblGrid')
col_widths_dxa = [410, 1390, 977, 1275, 936, 3243, 852]
for w in col_widths_dxa:
    col = OxmlElement('w:gridCol')
    col.set(qn('w:w'), str(w))
    tblGrid3.append(col)
phase_table._tbl.insert(0, tblGrid3)

# Header row
headers = ['\u2116', 'Phase name', 'Chemical formula', 'Crystal system', 'Space group', 'Lattice parameters, \u00c5', 'DB card number']
for j, h in enumerate(headers):
    cell = phase_table.cell(0, j)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    # Set cell width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(col_widths_dxa[j]))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# Data row (empty)
for j in range(7):
    cell = phase_table.cell(1, j)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('')
    r.font.size = Pt(9)
    r.font.name = 'Times New Roman'
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(col_widths_dxa[j]))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

# --- Figure placeholder ---
p = new_para(outer_cell, '', size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=3)
r = p.add_run('[ {{МЕСТО_ДЛЯ_ДИФРАКТОГРАММЫ}} ]')
r.italic = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

p = new_para(outer_cell, 'Рисунок 1 \u2014 Дифрактограмма образца {{ОБОЗНАЧЕНИЕ}}',
             bold=False, size=11, space_after=12)

# --- Conclusion ---
p = new_para(outer_cell, 'Заключение:', bold=True, space_before=6)
p.paragraph_format.line_spacing = 1.058

p = new_para(outer_cell, '\u2022 [карбиды]', size=10, left_indent=1.0, space_before=1)
p = new_para(outer_cell, '\u2022 [оксиды]', size=10, left_indent=1.0, space_before=1)
p = new_para(outer_cell, '\u2022 [прочие фазы]', size=10, left_indent=1.0, space_before=1)

p = new_para(outer_cell, '', size=10, left_indent=1.0, space_before=3)
r = p.add_run('{{ТЕКСТ_ЗАКЛЮЧЕНИЯ}}')
r.font.size = Pt(10)
r.font.name = 'Times New Roman'

p = new_para(outer_cell, '', size=10, left_indent=1.0, space_before=3)
r = p.add_run('Примечание: {{ПРИМЕЧАНИЕ_ПРИ_НАЛИЧИИ}}')
r.italic = True
r.font.size = Pt(10)
r.font.name = 'Times New Roman'

# --- Signature Table (3 columns, no borders) ---
# Add spacing
outer_cell.add_paragraph()
outer_cell.add_paragraph()

sig_table = outer_cell.add_table(rows=1, cols=3)
sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT

# Remove all borders
tblPr4 = sig_table._tbl.tblPr
tblBorders4 = parse_xml(
    '<w:tblBorders %s>'
    '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>' % nsdecls('w'))
tblPr4.append(tblBorders4)

# Set column widths
tblGrid4 = OxmlElement('w:tblGrid')
sig_widths = [3104, 2657, 3332]
for w in sig_widths:
    col = OxmlElement('w:gridCol')
    col.set(qn('w:w'), str(w))
    tblGrid4.append(col)
sig_table._tbl.insert(0, tblGrid4)

# Column 1: Position title
sc1 = sig_table.cell(0, 0)
set_cell_noborder(sc1)
set_cell_vertical_center(sc1)
p = sc1.paragraphs[0]
r = p.add_run('{{ДОЛЖНОСТЬ}}')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# Column 2: Signature image placeholder
sc2 = sig_table.cell(0, 1)
set_cell_noborder(sc2)
set_cell_vertical_center(sc2)
p = sc2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[ {{МЕСТО_ДЛЯ_ПОДПИСИ}} ]')
r.italic = True
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# Column 3: Name
sc3 = sig_table.cell(0, 2)
set_cell_noborder(sc3)
set_cell_vertical_center(sc3)
p = sc3.paragraphs[0]
r = p.add_run('{{ФАМИЛИЯ_ИНИЦИАЛЫ}}')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

# ============================================================
# FOOTER
# ============================================================
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT

r = fp.add_run('ПРОТОКОЛ ФИЗИКО-ХИМИЧЕСКОГО ФАЗОВОГО АНАЛИЗА № ')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

r = fp.add_run('{{НОМЕР_ПРОТОКОЛА}}')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

r = fp.add_run('   Лист ')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

add_field(fp, 'PAGE', bold=True, size=11)

r = fp.add_run(' из ')
r.font.size = Pt(11)
r.font.name = 'Times New Roman'

add_field(fp, 'NUMPAGES', bold=True, size=11)

# ============================================================
# SAVE
# ============================================================
temp_path = r'D:\temp_template_output.docx'
doc.save(temp_path)
print('Temp saved successfully')

# Move to target location
target_dir = r'D:\ИИ\Git Projetcs\Nushensia\WorkFlow\Паспорта\Шаблоны'
target_name = 'Шаблон_протокола_ФА_Open.docx'
target_path = os.path.join(target_dir, target_name)

# Remove lock file if exists
for f in os.listdir(target_dir):
    if f.startswith('~$') and target_name in f:
        os.remove(os.path.join(target_dir, f))

if os.path.exists(target_path):
    os.remove(target_path)
os.replace(temp_path, target_path)
print('Final file saved:', target_path)
print('Size:', os.path.getsize(target_path), 'bytes')
