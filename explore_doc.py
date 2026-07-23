import docx
doc = docx.Document(r'Отчеты/OpenCode/ОТЧЕТ_П-673_с_правками.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')
for i, p in enumerate(doc.paragraphs):
    text = p.text[:120] if p.text else ''
    style = p.style.name if p.style else 'None'
    fmt = p.paragraph_format
    left = fmt.left_indent
    align = fmt.alignment
    if text or style != 'Normal':
        print(f'  p[{i}]: style="{style}" align={align} left={left} "{text}"')
